from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_proposal():
    doc = Document()

    # Set margins to maximize space (Narrow margins)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Helper to set font
    def set_font(run, font_name='Malgun Gothic', size=10, bold=False, color=None):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = color

    # 1. Header / Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Snowball: 내부통제(ICFR) 혁신 솔루션")
    set_font(run, size=24, bold=True, color=RGBColor(44, 82, 101)) # #2c5265

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run("ITGC 전문 인력 부재와 수동 관리의 비효율성을 동시에 해결합니다.")
    set_font(run, size=12, color=RGBColor(100, 116, 139)) # Slate-500
    
    doc.add_paragraph() # Spacer

    # 2. Problem (Pain Points)
    heading_problem = doc.add_paragraph()
    run = heading_problem.add_run("1. 현재 내부통제 운영의 문제점 (Pain Points)")
    set_font(run, size=14, bold=True, color=RGBColor(44, 82, 101))

    problem_table = doc.add_table(rows=1, cols=2)
    problem_table.autofit = True
    
    # Cell 1
    cell1 = problem_table.cell(0, 0)
    p1 = cell1.paragraphs[0]
    run = p1.add_run("👤 IT 전문 인력 부재")
    set_font(run, size=11, bold=True)
    p1.add_run("\n내부통제팀 내 ITGC 전담 전문가 채용/유지 곤란\n높은 인건비와 이직률로 인한 업무 공백 발생")
    
    # Cell 2
    cell2 = problem_table.cell(0, 1)
    p2 = cell2.paragraphs[0]
    run = p2.add_run("⏰ 비효율적 수동 관리")
    set_font(run, size=11, bold=True)
    p2.add_run("\n엑셀 기반 수동 관리로 인한 버전 관리 오류\n단순 반복 업무로 핵심 리스크 관리 소홀")

    doc.add_paragraph() # Spacer

    # 3. Solution (Snowball)
    heading_solution = doc.add_paragraph()
    run = heading_solution.add_run("2. Snowball의 두 가지 솔루션")
    set_font(run, size=14, bold=True, color=RGBColor(44, 82, 101))

    run = doc.add_paragraph().add_run("Snowball은 PA 서비스(인력)와 시스템(솔루션)을 결합한 유일한 하이브리드 파트너입니다.")
    set_font(run, size=10)

    # Track A & B Table
    sol_table = doc.add_table(rows=1, cols=2)
    sol_table.style = 'Table Grid'
    
    # Track A
    c1 = sol_table.cell(0, 0)
    p = c1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nTrack A: PA 서비스 (전문 인력)")
    set_font(run, size=12, bold=True, color=RGBColor(37, 99, 235)) # Blue
    
    p = c1.add_paragraph()
    run = p.add_run("ITGC 업무 전 과정 상시 대행")
    set_font(run, size=10, bold=True)
    
    items_a = [
        "RCM 설계 및 표준화 (최신화, 설계평가)",
        "운영평가 대행 (샘플링, 증적 검증)",
        "보고서/조서 산출 및 감사 대응 지원"
    ]
    for item in items_a:
        p = c1.add_paragraph(item, style='List Bullet')
        set_font(p.runs[0], size=9)

    # Track B
    c2 = sol_table.cell(0, 1)
    p = c2.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nTrack B: Snowball 시스템")
    set_font(run, size=12, bold=True, color=RGBColor(5, 150, 105)) # Emerald
    
    p = c2.add_paragraph()
    run = p.add_run("내부통제 통합 관리 플랫폼")
    set_font(run, size=10, bold=True)

    items_b = [
        "Cloud(SaaS) 및 On-Premise(구축형) 지원",
        "ELC, TLC, ITGC 통합 대시보드",
        "평가 자동화 및 엑셀 조서 자동 생성"
    ]
    for item in items_b:
        p = c2.add_paragraph(item, style='List Bullet')
        set_font(p.runs[0], size=9)

    doc.add_paragraph() # Spacer

    # 4. Benefits
    heading_benefit = doc.add_paragraph()
    run = heading_benefit.add_run("3. 도입 효과")
    set_font(run, size=14, bold=True, color=RGBColor(44, 82, 101))

    benefits = [
        ("🛡️ 감사 대응력 강화", "외부 감사 질의 및 이슈 신속 해결"),
        ("🚀 업무 효율화", "반복 수동 업무 제거 및 자동화"),
        ("📊 운영 가시성 확보", "실시간 대시보드로 통제 현황 파악"),
        ("🎯 리스크 관리", "체계적 결함 추적 및 시정 조치")
    ]
    
    ben_table = doc.add_table(rows=2, cols=2)
    for i, (title, desc) in enumerate(benefits):
        row = i // 2
        col = i % 2
        cell = ben_table.cell(row, col)
        p = cell.paragraphs[0]
        run = p.add_run(title)
        set_font(run, size=11, bold=True)
        p.add_run(f"\n{desc}")

    doc.add_paragraph() # Spacer

    # 5. Contact / CTA
    cta_para = doc.add_paragraph()
    cta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = cta_para.add_run("지금 바로 Snowball 전문가와 상담하세요")
    set_font(run, size=16, bold=True, color=RGBColor(44, 82, 101))
    
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact_para.add_run("📧 snowball1566@gmail.com    🌐 www.snowball1566.com")
    set_font(run, size=12)

    # Save
    filename = 'snowball_proposal_1page.docx'
    doc.save(filename)
    print(f"Successfully created {filename}")

if __name__ == "__main__":
    create_proposal()
